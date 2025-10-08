import requests
import json
import os
import io
from docx import Document

BASE_URL = 'http://localhost:5000/api'

def create_sample_docx():
    """Create a sample DOCX file for testing"""
    doc = Document()
    doc.add_heading('Purchase Order Process', 0)
    
    # Add process steps
    steps = [
        "1. Purchase Request Creation",
        "- Employee fills out purchase request form",
        "- Includes item details, quantity, and cost estimates",
        "2. Department Manager Review",
        "- Reviews purchase request details",
        "- Checks budget availability",
        "3. Finance Department Review",
        "- Validates cost estimates",
        "- Checks vendor credentials",
        "4. Purchase Order Generation",
        "- Creates official purchase order",
        "- Records in ERP system",
        "5. Vendor Communication",
        "- Sends PO to vendor",
        "- Confirms delivery timeline",
        "6. Documentation and Closure",
        "- Archives all documents",
        "- Updates procurement records"
    ]
    
    for step in steps:
        doc.add_paragraph(step)
    
    # Save to memory
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io

def test_document_upload():
    """Test document upload endpoint"""
    print("\nCreating test DOCX document...")
    docx_io = create_sample_docx()
    
    print("Uploading document...")
    files = {
        'file': ('purchase_order_process.docx', docx_io, 
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    }
    
    response = requests.post(f'{BASE_URL}/documents', files=files)
    
    print('\nDocument Upload Test:')
    print(f'Status Code: {response.status_code}')
    try:
        print(f'Response: {response.json()}')
        return response.json().get('document', {}).get('id')
    except:
        print(f'Response: {response.text}')
        return None

def test_workflow_extraction(document_id):
    """Test workflow extraction endpoint"""
    if not document_id:
        print("Error: No document ID provided for workflow extraction")
        return
    
    print("\nExtracting workflow...")
    response = requests.post(f'{BASE_URL}/workflows/extract/{document_id}')
    
    print('\nWorkflow Extraction Test:')
    print(f'Status Code: {response.status_code}')
    try:
        print(f'Response: {response.json()}')
        return response.json().get('workflow', {}).get('id')
    except:
        print(f'Response: {response.text}')
        return None

def test_workflow_analysis(workflow_id):
    """Test workflow analysis endpoint"""
    if not workflow_id:
        print("Error: No workflow ID provided for analysis")
        return
    
    print("\nAnalyzing workflow...")
    response = requests.post(f'{BASE_URL}/workflows/{workflow_id}/analyze')
    
    print('\nWorkflow Analysis Test:')
    print(f'Status Code: {response.status_code}')
    try:
        print(f'Response: {response.json()}')
        return True
    except:
        print(f'Response: {response.text}')
        return False

def test_workflow_optimization(workflow_id):
    """Test workflow optimization endpoint"""
    if not workflow_id:
        print("Error: No workflow ID provided for optimization")
        return
    
    print("\nOptimizing workflow...")
    response = requests.post(f'{BASE_URL}/workflows/{workflow_id}/optimize')
    
    print('\nWorkflow Optimization Test:')
    print(f'Status Code: {response.status_code}')
    try:
        print(f'Response: {response.json()}')
    except:
        print(f'Response: {response.text}')

def test_erp_process_creation():
    """Test ERP process creation with steps"""
    print("\nCreating ERP process...")
    data = {
        "process_name": "Purchase Order Process",
        "erp_system": "SAP",
        "steps": [
            {
                "step_number": 1,
                "step_description": "Purchase Request Creation",
                "approver_role": "Employee",
                "approver_name": "John Smith"
            },
            {
                "step_number": 2,
                "step_description": "Department Manager Review",
                "approver_role": "Department Manager",
                "approver_name": "Jane Wilson"
            },
            {
                "step_number": 3,
                "step_description": "Finance Department Review",
                "approver_role": "Finance Manager",
                "approver_name": "Mike Johnson"
            },
            {
                "step_number": 4,
                "step_description": "Purchase Order Generation",
                "approver_role": "Procurement Officer",
                "approver_name": "Sarah Davis"
            }
        ]
    }
    
    response = requests.post(
        f'{BASE_URL}/erp-processes',
        json=data,
        headers={'Content-Type': 'application/json'}
    )
    
    print('\nERP Process Creation Test:')
    print(f'Status Code: {response.status_code}')
    try:
        print(f'Response: {response.json()}')
        return response.json().get('process', {}).get('id')
    except:
        print(f'Response: {response.text}')
        return None

def test_erp_process_retrieval(process_id):
    """Test retrieving an ERP process"""
    if not process_id:
        print("Error: No process ID provided for retrieval")
        return
    
    print("\nRetrieving ERP process...")
    response = requests.get(f'{BASE_URL}/erp-processes/{process_id}')
    
    print('\nERP Process Retrieval Test:')
    print(f'Status Code: {response.status_code}')
    try:
        print(f'Response: {response.json()}')
    except:
        print(f'Response: {response.text}')

if __name__ == '__main__':
    print("Testing AI-Powered BPM System API...")
    print("=====================================")
    
    # Test document upload and workflow extraction
    print("\n1. Testing Document Processing...")
    document_id = test_document_upload()
    
    if document_id:
        print("\n2. Testing Workflow Features...")
        workflow_id = test_workflow_extraction(document_id)
        
        if workflow_id:
            # Test workflow analysis
            print("\n3. Testing Workflow Analysis...")
            if test_workflow_analysis(workflow_id):
                # Test workflow optimization
                print("\n4. Testing Workflow Optimization...")
                test_workflow_optimization(workflow_id)
    
    # Test ERP process management
    print("\n5. Testing ERP Process Management...")
    process_id = test_erp_process_creation()
    
    if process_id:
        print("\n6. Testing ERP Process Retrieval...")
        test_erp_process_retrieval(process_id)
